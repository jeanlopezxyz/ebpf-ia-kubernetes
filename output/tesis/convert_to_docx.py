#!/usr/bin/env python3
"""
Convert the thesis markdown to a properly formatted DOCX document.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import re

def create_thesis_docx():
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)
    
    # Define styles
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Times New Roman'
    title_font.size = Pt(16)
    title_font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Heading 1 style
    heading1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_font = heading1_style.font
    heading1_font.name = 'Times New Roman'
    heading1_font.size = Pt(14)
    heading1_font.bold = True
    heading1_style.paragraph_format.space_before = Pt(18)
    heading1_style.paragraph_format.space_after = Pt(6)
    
    # Heading 2 style
    heading2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_font = heading2_style.font
    heading2_font.name = 'Times New Roman'
    heading2_font.size = Pt(12)
    heading2_font.bold = True
    heading2_style.paragraph_format.space_before = Pt(12)
    heading2_style.paragraph_format.space_after = Pt(6)
    
    # Normal style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(11)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.15
    
    # Read the markdown file
    with open('/home/jeanlopez/Documents/personal/projects/ebpf-ia-kubernetes/output/tesis_10_paginas_ebpf_ia.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split content into lines
    lines = content.split('\n')
    
    in_table = False
    table = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Handle UTP title
        if line.startswith('# ') and 'Universidad Tecnológica del Perú' in line:
            title = line[2:].strip()
            p = doc.add_paragraph(title, style='TitleStyle')
            continue
            
        # Handle main thesis title
        elif line.startswith('**"IMPLEMENTACIÓN DE UNA PLATAFORMA'):
            title = line.replace('**"', '').replace('"**', '')
            p = doc.add_paragraph(title, style='TitleStyle')
            continue
            
        # Handle section headers
        elif line.startswith('## '):
            header = line[3:].strip()
            p = doc.add_paragraph(header, style='CustomHeading1')
            continue
            
        # Handle subsection headers
        elif line.startswith('### '):
            subheader = line[4:].strip()
            p = doc.add_paragraph(subheader, style='CustomHeading2')
            continue
            
        # Handle UTP faculty and program info
        elif line.startswith('**Facultad de Ingeniería**') or line.startswith('**Carrera de Ingeniería de Telecomunicaciones**'):
            p = doc.add_paragraph(line.replace('**', ''), style='Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        # Handle UTP author name
        elif line.startswith('**Jean López**'):
            p = doc.add_paragraph(line.replace('**', ''), style='Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.bold = True
            continue
            
        # Handle UTP degree statement
        elif 'para optar el Título Profesional de Ingeniero de Telecomunicaciones' in line:
            p = doc.add_paragraph(line.replace('*', ''), style='Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.italic = True
            continue
            
        # Handle UTP advisor
        elif line.startswith('**Asesor:**'):
            p = doc.add_paragraph(line.replace('**', ''), style='Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        # Handle UTP location and date
        elif line.startswith('**Lima – Perú**') or line.startswith('**Septiembre 2025**'):
            p = doc.add_paragraph(line.replace('**', ''), style='Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        # Handle horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph()
            continue
            
        # Handle table headers
        elif '|' in line and ('Tipo de Amenaza' in line or 'Precisión' in line):
            if not in_table:
                table = doc.add_table(rows=0, cols=5)
                table.style = 'Table Grid'
                in_table = True
                
                # Add header row
                header_row = table.add_row()
                cells = line.split('|')[1:-1]  # Remove empty first and last
                for i, cell in enumerate(cells):
                    if i < len(header_row.cells):
                        header_row.cells[i].text = cell.strip()
                        # Make header bold
                        for paragraph in header_row.cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
            continue
            
        # Handle table separator
        elif '|---' in line:
            continue
            
        # Handle table data rows
        elif in_table and '|' in line and line.count('|') >= 4:
            data_row = table.add_row()
            cells = line.split('|')[1:-1]  # Remove empty first and last
            for i, cell in enumerate(cells):
                if i < len(data_row.cells):
                    data_row.cells[i].text = cell.strip()
            continue
            
        # End table if we're not in a table line anymore
        elif in_table and '|' not in line:
            in_table = False
            table = None
            
        # Handle code blocks
        elif line.startswith('```'):
            continue
            
        # Handle bullet points
        elif line.startswith('- **') or line.startswith('1. **') or line.startswith('2. **'):
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            p = doc.add_paragraph(text, style='Normal')
            continue
            
        # Handle regular bold text
        elif '**' in line:
            p = doc.add_paragraph(style='Normal')
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    # Regular text
                    if part:
                        p.add_run(part)
                else:
                    # Bold text
                    if part:
                        run = p.add_run(part)
                        run.font.bold = True
            continue
            
        # Handle Keywords line
        elif line.startswith('**Palabras clave:**') or line.startswith('**Keywords:**'):
            p = doc.add_paragraph(line.replace('**', ''), style='Normal')
            continue
            
        # Handle regular paragraphs
        else:
            if line and not line.startswith('#'):
                p = doc.add_paragraph(line, style='Normal')
                continue
    
    # Add page break before references if needed
    last_para = doc.paragraphs[-1]
    if 'REFERENCIAS' in last_para.text:
        last_para.add_run().add_break()
    
    # Save the document
    output_path = '/home/jeanlopez/Documents/personal/projects/ebpf-ia-kubernetes/output/tesis_10_paginas_UTP_ebpf_ia.docx'
    doc.save(output_path)
    print(f"DOCX document saved to: {output_path}")
    
    # Print document statistics
    print(f"Document contains {len(doc.paragraphs)} paragraphs")
    
    return output_path

if __name__ == "__main__":
    create_thesis_docx()